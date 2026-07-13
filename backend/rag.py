import os
import uuid
from pathlib import Path

from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

load_dotenv()

CHROMA_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "chroma_db")
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "uploads")

Path(UPLOAD_DIR).mkdir(exist_ok=True)

CHUNK_SIZE = 500
CHUNK_OVERLAP = 100
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

_text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", "。", ".", " ", ""],
)

_embedding_fn = None
_vectorstore = None


class ChromadbEmbeddingAdapter(Embeddings):
    """Wrap chromadb's native embedding function to match langchain's Embeddings interface."""

    def __init__(self, ef):
        self._ef = ef

    def embed_documents(self, texts):
        return self._ef(texts).tolist()

    def embed_query(self, text):
        return self._ef(text).tolist()


def _get_embedding_fn():
    global _embedding_fn
    if _embedding_fn is None:
        import chromadb.utils.embedding_functions as ef
        _embedding_fn = ChromadbEmbeddingAdapter(ef.DefaultEmbeddingFunction())
    return _embedding_fn


def _get_vectorstore():
    global _vectorstore
    if _vectorstore is None:
        _vectorstore = Chroma(
            persist_directory=CHROMA_DIR,
            embedding_function=_get_embedding_fn(),
            collection_name="meiken_kb",
        )
    return _vectorstore


def extract_text(filepath: str, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        reader = PdfReader(filepath)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext in (".txt", ".md"):
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(filepath)
        return "\n\n".join(p.text for p in doc.paragraphs)
    else:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def ingest_document(filepath: str, filename: str, file_id: str, scope: str = "temp"):
    """Extract text, split into chunks, and store in Chroma vector DB."""
    text = extract_text(filepath, filename)
    if not text.strip():
        return 0

    chunks = _text_splitter.split_text(text)
    docs = [
        Document(
            page_content=chunk,
            metadata={"file_id": file_id, "filename": filename, "scope": scope},
        )
        for chunk in chunks
    ]
    _get_vectorstore().add_documents(docs)
    return len(chunks)


def search_relevant(query: str, scope: str = None, k: int = 4, file_ids: list = None):
    """Search for relevant document chunks. If scope is set, filter by it. If file_ids is set, filter by file_id."""
    vs = _get_vectorstore()
    if file_ids:
        filter_dict = {"file_id": {"$in": file_ids}}
    elif scope:
        filter_dict = {"scope": scope}
    else:
        filter_dict = None
    results = vs.similarity_search_with_relevance_scores(query, k=k, filter=filter_dict)
    return [
        {
            "content": doc.page_content,
            "filename": doc.metadata.get("filename", ""),
            "file_id": doc.metadata.get("file_id", ""),
            "score": round(score, 3),
        }
        for doc, score in results
    ]


def delete_document(file_id: str):
    """Remove all chunks belonging to a file from the vector store."""
    _get_vectorstore().delete(where={"file_id": file_id})


def save_upload_file(file_content: bytes, filename: str) -> str:
    """Save uploaded file to disk, return filepath."""
    file_id = uuid.uuid4().hex[:12]
    ext = Path(filename).suffix
    filepath = os.path.join(UPLOAD_DIR, f"{file_id}{ext}")
    with open(filepath, "wb") as f:
        f.write(file_content)
    return filepath, file_id


def cleanup_upload_file(filepath: str):
    """Remove uploaded file from disk."""
    try:
        os.remove(filepath)
    except OSError:
        pass
